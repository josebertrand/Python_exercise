import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
import comtypes.client
import subprocess  # Para ejecutar comandos de sistema
import shutil  # Para mover archivos

# Rutas de archivos
excel_path = "C:\\Users\\José Bertrand\\Desktop\\prueba pythom\\Arch.xlsx"
template_path = "C:\\Users\\José Bertrand\\Desktop\\prueba pythom\\plantilla.docx"
output_dir = "C:\\Users\\José Bertrand\\Desktop\\prueba pythom\\Documentos Generados"
pdf_output_dir = "C:\\Users\\José Bertrand\\Desktop\\prueba pythom\\Documentos Generados PDF"
protected_pdf_dir = "C:\\Users\\José Bertrand\\Desktop\\prueba pythom\\Documentos Protegidos"  # Nueva carpeta para los PDFs protegidos

# Crear el directorio de salida si no existe
os.makedirs(output_dir, exist_ok=True)
os.makedirs(pdf_output_dir, exist_ok=True)
os.makedirs(protected_pdf_dir, exist_ok=True)  # Crear el directorio para los PDFs protegidos

# Cargar datos desde Excel
try:
    df = pd.read_excel(excel_path)
    df.columns = df.columns.str.strip()  # Eliminar espacios en los nombres de las columnas
except Exception as e:
    print(f"Error al cargar el archivo Excel: {e}")
    exit()

# Mapa de meses en español
meses_espanol = {
    "January": "enero", "February": "febrero", "March": "marzo",
    "April": "abril", "May": "mayo", "June": "junio",
    "July": "julio", "August": "agosto", "September": "septiembre",
    "October": "octubre", "November": "noviembre", "December": "diciembre"
}

# Función para proteger el PDF con permisos (solo impresión)
def proteger_pdf_con_permisos(pdf_path, output_path):
    try:
        # Ejecutar el comando pdftk para proteger el PDF
        # Esto asume que pdftk está en el PATH
        command = [
            "pdftk", pdf_path, "output", output_path, 
            "owner_pw", "Zamo10*",   # Contraseña del propietario
            "allow", "printing"         # Solo permite impresión, no edición ni copia
        ]
        subprocess.run(command, check=True)  # Ejecuta el comando
        print(f"El archivo PDF protegido se ha guardado en: {output_path}")
    except Exception as e:
        print(f"Error al proteger el PDF: {e}")

# Generar documentos a partir de los datos en Excel
documentos_generados = 0
pdfs_protegidos = 0  # Contador de PDFs protegidos

for index, row in df.iterrows():
    try:
        # Extraer datos de cada fila
        nombre = str(row['Nombres']).strip()
        codigo = str(row['Código']).strip()
        departamento = str(row['departamento']).strip()
        titulo = str(row['Título']).strip()

        if not nombre or not codigo or not departamento or not titulo:
            print(f"Faltan datos en la fila {index + 1}. Saltando...")
            continue

        # Formato de la fecha en español
        fecha = datetime.now().strftime("%A, %d de %B de %Y")
        fecha = fecha.replace("Monday", "lunes").replace("Tuesday", "martes").replace(
            "Wednesday", "miércoles").replace("Thursday", "jueves").replace(
            "Friday", "viernes").replace("Saturday", "sábado").replace(
            "Sunday", "domingo")
        
        for mes_en, mes_es in meses_espanol.items():
            fecha = fecha.replace(mes_en, mes_es)

        # Abrir documento de plantilla
        doc = Document(template_path)

        # Insertar información en el documento
        for paragraph in doc.paragraphs:
            if "{{nombre}}" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run("Que el estudiante ")
                run.font.size = Pt(12)

                run_nombre = paragraph.add_run(nombre)
                run_nombre.bold = True
                run_nombre.font.size = Pt(12)

                run = paragraph.add_run(" con código No. ")
                run.font.size = Pt(12)

                run_codigo = paragraph.add_run(codigo)
                run_codigo.bold = True
                run_codigo.font.size = Pt(12)

                run = paragraph.add_run(" del ")
                run.font.size = Pt(12)

                run_departamento = paragraph.add_run(departamento)
                run_departamento.bold = True
                run_departamento.font.size = Pt(12)

                run = paragraph.add_run(" hizo entrega de PEG titulada ")
                run.font.size = Pt(12)

                run_titulo = paragraph.add_run(titulo)
                run_titulo.bold = True
                run_titulo.font.size = Pt(12)

                run = paragraph.add_run(" destete, acompañado de los siguientes Documentos: ")
                run.font.size = Pt(12)
            elif "{{fecha}}" in paragraph.text:
                # Limpiar el marcador y añadir la fecha
                paragraph.clear()
                run = paragraph.add_run("Dada en Zamorano el ")
                run.font.size = Pt(12)

                run_fecha = paragraph.add_run(fecha)
                run_fecha.bold = True
                run_fecha.font.size = Pt(12)

        # Guardar el documento generado en formato Word
        doc_output_path = os.path.join(output_dir, f"{nombre.replace(' ', '_')}.docx")
        doc.save(doc_output_path)
        documentos_generados += 1
        print(f"Documento generado en Word: {doc_output_path}")

        # Convertir el archivo DOCX a PDF usando Microsoft Word
        pdf_output_path = os.path.join(pdf_output_dir, f"{nombre.replace(' ', '_')}.pdf")
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(doc_output_path)
        doc.SaveAs(pdf_output_path, FileFormat=17)  # 17 es el formato PDF
        doc.Close()
        word.Quit()
        print(f"Documento convertido a PDF: {pdf_output_path}")

        # Proteger el PDF generado con permisos
        pdf_protegido_output_path = os.path.join(pdf_output_dir, f"{nombre.replace(' ', '_')}_Protegido.pdf")
        proteger_pdf_con_permisos(pdf_output_path, pdf_protegido_output_path)

        # Mover el PDF protegido a la carpeta correspondiente
        protected_pdf_final_path = os.path.join(protected_pdf_dir, f"{nombre.replace(' ', '_')}_Protegido.pdf")
        shutil.move(pdf_protegido_output_path, protected_pdf_final_path)
        pdfs_protegidos += 1
        print(f"PDF protegido movido a: {protected_pdf_final_path}")

    except Exception as e:
        print(f"Error al generar el documento para {nombre}: {e}")

print(f"Total de documentos generados: {documentos_generados}")
print(f"Total de PDFs protegidos: {pdfs_protegidos}")
